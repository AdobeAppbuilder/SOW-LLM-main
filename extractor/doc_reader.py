import base64
import hashlib
import io
from typing import Tuple
from docx import Document
from pypdf import PdfReader


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def extract_text_from_bytes(filename: str, file_bytes_base64: str) -> Tuple[str, str]:
    """Return (extracted_text, sha256_checksum)."""
    raw = base64.b64decode(file_bytes_base64)
    checksum = _sha256(raw)

    lower = filename.lower()
    if lower.endswith('.docx'):
        return _extract_docx(raw), checksum
    if lower.endswith('.pdf'):
        return _extract_pdf(raw), checksum

    # Best-effort: treat as utf-8 text
    try:
        return raw.decode('utf-8', errors='ignore'), checksum
    except Exception as e:
        raise ValueError(f"Unsupported file type for {filename}") from e


def _extract_docx(raw: bytes) -> str:
    bio = io.BytesIO(raw)
    doc = Document(bio)
    parts = []

    for p in doc.paragraphs:
        t = (p.text or '').strip()
        if t:
            parts.append(t)

    # tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(' | '.join(cells))

    return ''.join(parts)


def _extract_pdf(raw: bytes) -> str:
    bio = io.BytesIO(raw)
    reader = PdfReader(bio)
    parts = []
    for page in reader.pages:
        txt = (page.extract_text() or '').strip()
        if txt:
            parts.append(txt)
    return ''.join(parts)
